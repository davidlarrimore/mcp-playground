import logging
import os
from pathlib import Path
from typing import List, Optional
from fastmcp import FastMCP
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

LOG_LEVEL = os.getenv("LOG_LEVEL", "INFO").upper()
logging.basicConfig(level=LOG_LEVEL, format="%(asctime)s [%(levelname)s] %(name)s - %(message)s")
logger = logging.getLogger("word-mcp")

WORKSPACE = Path(os.getenv("WORKSPACE", "/workspace")).resolve()

mcp = FastMCP("word")

@mcp.tool
async def save_uploaded_file(
    filename: str,
    content_base64: str
) -> dict:
    """
    Save an uploaded file (provided as base64) to the workspace.

    This is a convenience tool for when users upload files via chat.
    After saving, the file will be accessible to all workspace tools.

    IMPORTANT: Preserve the original file extension! If user uploads "data.xlsx",
    save it as "data.xlsx" (or similar .xlsx name), NOT as "data.txt".

    Args:
        filename: Name to save the file as - MUST keep original file extension
        content_base64: Base64-encoded file content from uploaded file
    """
    import base64

    try:
        file_data = base64.b64decode(content_base64)
        output_path = WORKSPACE / filename
        output_path.write_bytes(file_data)

        logger.info(f"Saved uploaded file: {output_path}")

        return {
            "path": str(output_path.relative_to(WORKSPACE)),
            "absolute_path": str(output_path),
            "size": output_path.stat().st_size,
            "message": f"File saved successfully as '{filename}'"
        }
    except Exception as e:
        error_msg = f"Error saving file: {str(e)}"
        logger.error(error_msg)
        return {"error": error_msg}

@mcp.tool
async def create_word_report(
    title: str,
    sections: List[dict],
    output_filename: str,
    subtitle: Optional[str] = None
) -> dict:
    """
    Create a professional Word document.

    Args:
        title: Report title
        sections: List of sections with structure:
            [{
                "heading": "Executive Summary",
                "level": 1,
                "content": "This week's performance shows...",
                "table": {
                    "columns": ["Region", "Revenue"],
                    "rows": [{"Region": "East", "Revenue": "50K"}]
                },
                "bullet_points": ["Point 1", "Point 2"]
            }]
        output_filename: Name of output file
        subtitle: Optional subtitle (e.g., "Week Ending Dec 1, 2024")
    """
    doc = Document()

    # Title
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if subtitle:
        subtitle_para = doc.add_paragraph(subtitle)
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_para.runs[0].font.size = Pt(14)
        subtitle_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

    # Sections
    for section in sections:
        level = section.get("level", 1)
        doc.add_heading(section["heading"], level)

        if "content" in section:
            doc.add_paragraph(section["content"])

        if "bullet_points" in section:
            for point in section["bullet_points"]:
                doc.add_paragraph(point, style='List Bullet')

        if "table" in section:
            table_data = section["table"]
            table = doc.add_table(
                rows=len(table_data["rows"]) + 1,
                cols=len(table_data["columns"])
            )
            table.style = 'Light Grid Accent 1'

            # Headers
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(table_data["columns"]):
                hdr_cells[i].text = col
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True

            # Data rows
            for row_idx, row in enumerate(table_data["rows"], 1):
                row_cells = table.rows[row_idx].cells
                for col_idx, col_name in enumerate(table_data["columns"]):
                    row_cells[col_idx].text = str(row.get(col_name, ""))

    output_path = WORKSPACE / output_filename
    doc.save(output_path)

    logger.info(f"Created Word document: {output_path}")

    return {
        "path": str(output_path.relative_to(WORKSPACE)),
        "absolute_path": str(output_path),
        "size": output_path.stat().st_size,
        "sections": len(sections)
    }

if __name__ == "__main__":
    mcp.run()
