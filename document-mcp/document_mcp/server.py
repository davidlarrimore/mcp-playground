import logging
import os
import sys
from pathlib import Path
from typing import List, Optional

from fastmcp import FastMCP
from openpyxl import Workbook, load_workbook
from openpyxl.chart import BarChart, LineChart, Reference
from openpyxl.styles import Alignment, Font, PatternFill
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from weasyprint import HTML, CSS

# Add parent directory to path for shared module
sys.path.insert(0, str(Path(__file__).parent.parent.parent))
from shared.download_urls import generate_signed_url  # noqa: E402

LOG_LEVEL = os.getenv("LOG_LEVEL", "INFO").upper()
logging.basicConfig(level=LOG_LEVEL, format="%(asctime)s [%(levelname)s] %(name)s - %(message)s")
logger = logging.getLogger("document-mcp")

WORKSPACE = Path(os.getenv("WORKSPACE", "/workspace")).resolve()

mcp = FastMCP(
    "document",
    instructions=(
        "Document helper for combining user-uploaded files and producing reports. "
        "Always read requested workspace files first. "
        "Data is required: when creating or merging workbooks you must supply real tabular data (columns + rows). "
        "Upload user files with save_uploaded_file (base64 content, keep extensions). "
        "To merge multiple Excel files, first upload them, then call combine_excel_files with the file paths. "
        "To build a workbook from data the model already has, call create_excel_workbook with non-empty columns and rows. "
        "Use create_word_report for formatted summaries and create_pdf_from_html for HTML-to-PDF."
    ),
)


def _with_download_metadata(output_path: Path, message: str, extra: Optional[dict] = None, download_filename: Optional[str] = None) -> dict:
    """
    Build a standard result payload with optional signed download URL metadata.
    """
    result = {
        "path": str(output_path.relative_to(WORKSPACE)),
        "absolute_path": str(output_path),
        "size": output_path.stat().st_size,
        "message": message
    }

    if extra:
        result.update(extra)

    try:
        download_info = generate_signed_url(str(output_path), filename=download_filename)
        expires_hours = download_info["expires_in"] // 3600
        result.update({
            "download_url": download_info["url"],
            "download_expires_at": download_info["expires_at"],
            "download_expires_in": download_info["expires_in"],
            "message": f"{message}\n\n📥 Download your file:\n{download_info['url']}\n\n⏰ Link expires in {expires_hours} hours"
        })
    except Exception as e:
        logger.error(f"Error generating download URL: {e}")

    return result


def _resolve_path(path_str: str) -> Path:
    """Resolve a workspace-relative path and prevent directory escape."""
    candidate = (WORKSPACE / path_str).resolve()
    if not candidate.is_relative_to(WORKSPACE):
        raise ValueError("Path must stay within the workspace")
    return candidate


def _load_excel_table(path: Path) -> tuple[list[str], list[dict]]:
    """Load the first sheet of an Excel file as headers plus row dictionaries."""
    wb = load_workbook(path, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return [], []

    headers = [str(h).strip() if h is not None else "" for h in rows[0]]
    headers = [h for h in headers if h]  # drop empty header cells
    data_rows = []
    for row in rows[1:]:
        row_dict = {}
        for idx, header in enumerate(headers):
            row_dict[header] = row[idx] if idx < len(row) else None
        data_rows.append(row_dict)
    return headers, data_rows


@mcp.tool
async def save_uploaded_file(
    filename: str,
    content_base64: str
) -> dict:
    """
    Save an uploaded file (provided as base64) to the workspace for downstream tools.

    IMPORTANT: Preserve the original file extension! If user uploads "data.xlsx",
    save it as "data.xlsx" (or similar .xlsx name), NOT as "data.txt".

    Args:
        filename: Name to save the file as - MUST keep original file extension
        content_base64: Base64-encoded file content from uploaded file
    """
    import base64

    try:
        file_data = base64.b64decode(content_base64)
        output_path = WORKSPACE / filename
        output_path.write_bytes(file_data)

        logger.info(f"Saved uploaded file: {output_path}")

        return {
            "path": str(output_path.relative_to(WORKSPACE)),
            "absolute_path": str(output_path),
            "size": output_path.stat().st_size,
            "message": f"File saved successfully as '{filename}'"
        }
    except Exception as e:
        error_msg = f"Error saving file: {str(e)}"
        logger.error(error_msg)
        return {"error": error_msg}


@mcp.tool
async def create_excel_workbook(
    sheets: List[dict],
    output_filename: str,
    add_charts: bool = True
) -> dict:
    """
    Create an Excel workbook with multiple sheets.

    Args:
        sheets: List of sheets with structure:
            [{
                "name": "Sheet1",
                "data": {
                    "columns": ["Region", "Revenue", "Orders"],
                    "rows": [
                        {"Region": "East", "Revenue": 50000, "Orders": 120},
                        {"Region": "West", "Revenue": 75000, "Orders": 180}
                    ]
                },
                "chart": {
                    "type": "bar",
                    "title": "Revenue by Region"
                }
            }]
        output_filename: Name of output file
        add_charts: Whether to add charts
    """
    wb = Workbook()
    wb.remove(wb.active)

    for sheet_def in sheets:
        ws = wb.create_sheet(title=sheet_def["name"])
        data = sheet_def["data"]
        columns = data.get("columns") or []
        rows = data.get("rows") or []
        if not columns:
            raise ValueError("create_excel_workbook requires non-empty 'columns'.")
        if not rows:
            raise ValueError("create_excel_workbook requires at least one row of data.")

        # Write headers
        for col_idx, col_name in enumerate(columns, 1):
            cell = ws.cell(row=1, column=col_idx, value=col_name)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="366092", fill_type="solid")
            cell.alignment = Alignment(horizontal="center")

        # Write data rows
        for row_idx, row_data in enumerate(rows, 2):
            for col_idx, col_name in enumerate(columns, 1):
                ws.cell(row=row_idx, column=col_idx, value=row_data.get(col_name))

        # Auto-adjust column widths
        for col in ws.columns:
            max_length = max(len(str(cell.value or "")) for cell in col)
            ws.column_dimensions[col[0].column_letter].width = max_length + 2

        # Add chart if requested
        if add_charts and "chart" in sheet_def:
            chart = BarChart() if sheet_def["chart"].get("type") == "bar" else LineChart()
            chart.title = sheet_def["chart"].get("title", "Chart")

            data_ref = Reference(ws, min_col=2, min_row=1, max_row=len(data["rows"]) + 1)
            cats = Reference(ws, min_col=1, min_row=2, max_row=len(data["rows"]) + 1)

            chart.add_data(data_ref, titles_from_data=True)
            chart.set_categories(cats)
            ws.add_chart(chart, "E5")

    output_path = WORKSPACE / output_filename
    wb.save(output_path)

    logger.info(f"Created Excel workbook: {output_path}")

    return _with_download_metadata(
        output_path,
        f"Excel workbook created successfully with {len(sheets)} sheet(s).",
        extra={"sheets": len(sheets)}
    )


@mcp.tool
async def create_word_report(
    title: str,
    sections: List[dict],
    output_filename: str,
    subtitle: Optional[str] = None
) -> dict:
    """
    Create a professional Word document.

    Args:
        title: Report title
        sections: List of sections with structure:
            [{
                "heading": "Executive Summary",
                "level": 1,
                "content": "This week's performance shows...",
                "table": {
                    "columns": ["Region", "Revenue"],
                    "rows": [{"Region": "East", "Revenue": "50K"}]
                },
                "bullet_points": ["Point 1", "Point 2"]
            }]
        output_filename: Name of output file
        subtitle: Optional subtitle (e.g., "Week Ending Dec 1, 2024")
    """
    doc = Document()

    # Title
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if subtitle:
        subtitle_para = doc.add_paragraph(subtitle)
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_para.runs[0].font.size = Pt(14)
        subtitle_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

    # Sections
    for section in sections:
        level = section.get("level", 1)
        doc.add_heading(section["heading"], level)

        if "content" in section:
            doc.add_paragraph(section["content"])

        if "bullet_points" in section:
            for point in section["bullet_points"]:
                doc.add_paragraph(point, style="List Bullet")

        if "table" in section:
            table_data = section["table"]
            table = doc.add_table(
                rows=len(table_data["rows"]) + 1,
                cols=len(table_data["columns"])
            )
            table.style = "Light Grid Accent 1"

            # Headers
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(table_data["columns"]):
                hdr_cells[i].text = col
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True

            # Data rows
            for row_idx, row in enumerate(table_data["rows"], 1):
                row_cells = table.rows[row_idx].cells
                for col_idx, col_name in enumerate(table_data["columns"]):
                    row_cells[col_idx].text = str(row.get(col_name, ""))

    output_path = WORKSPACE / output_filename
    doc.save(output_path)

    logger.info(f"Created Word document: {output_path}")

    return _with_download_metadata(
        output_path,
        f"Word document '{title}' created successfully with {len(sections)} section(s).",
        extra={"sections": len(sections)}
    )


@mcp.tool
async def create_pdf_from_html(
    html_content: str,
    output_filename: str,
    css: Optional[str] = None
) -> dict:
    """
    Create a PDF from HTML content.

    Args:
        html_content: HTML content or path to HTML file (relative to workspace)
        output_filename: Name of output PDF file
        css: Optional CSS stylesheet content for styling
    """
    output_path = WORKSPACE / output_filename

    # Treat html_content as inline HTML if it clearly contains markup or if checking
    # a file path would raise (e.g., because the string is long HTML and not a path).
    use_inline = any(token in html_content[:200].lower() for token in ["<html", "<!doctype", "<head", "<body"])
    if not use_inline:
        try:
            html_file = WORKSPACE / html_content
            if html_file.exists() and html_file.is_file():
                html = HTML(filename=str(html_file))
            else:
                html = HTML(string=html_content)
        except OSError:
            html = HTML(string=html_content)
    else:
        html = HTML(string=html_content)

    stylesheets = []
    if css:
        stylesheets.append(CSS(string=css))

    html.write_pdf(output_path, stylesheets=stylesheets)

    logger.info(f"Created PDF: {output_path}")

    return _with_download_metadata(
        output_path,
        f"PDF created successfully: {output_filename}"
    )


@mcp.tool
async def combine_excel_files(
    source_files: List[str],
    output_filename: str = "combined.xlsx",
    sheet_name: str = "Combined",
    include_source_column: bool = True,
) -> dict:
    """
    Combine the first sheet of multiple Excel files into a single sheet.

    Args:
        source_files: Workspace-relative Excel file paths to merge.
        output_filename: Name of the merged workbook to write.
        sheet_name: Name of the consolidated sheet.
        include_source_column: Append a column indicating which file each row came from.
    """
    if not source_files:
        raise ValueError("Provide at least one Excel file to combine.")

    columns: list[str] = []
    combined_rows: list[tuple[dict, str]] = []

    for path_str in source_files:
        file_path = _resolve_path(path_str)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {path_str}")
        if file_path.suffix.lower() not in {".xlsx", ".xlsm"}:
            raise ValueError(f"File must be an Excel workbook: {path_str}")

        headers, rows = _load_excel_table(file_path)
        if not headers:
            logger.warning("Skipped empty or headerless file: %s", path_str)
            continue

        for header in headers:
            if header not in columns:
                columns.append(header)

        for row in rows:
            combined_rows.append((row, file_path.name))

    if not combined_rows:
        raise ValueError("No row data found in the provided files.")

    source_col = "Source"
    if include_source_column and source_col in columns:
        source_col = "SourceFile"
    if include_source_column:
        columns.append(source_col)

    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name

    # Write headers
    for col_idx, col_name in enumerate(columns, 1):
        cell = ws.cell(row=1, column=col_idx, value=col_name)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="4BACC6", fill_type="solid")
        cell.alignment = Alignment(horizontal="center")

    # Write rows
    for row_idx, (row_data, source_name) in enumerate(combined_rows, start=2):
        for col_idx, col_name in enumerate(columns, 1):
            value = row_data.get(col_name)
            if include_source_column and col_name == source_col:
                value = source_name
            ws.cell(row=row_idx, column=col_idx, value=value)

    # Auto-fit column widths
    for col in ws.columns:
        max_length = max(len(str(cell.value or "")) for cell in col)
        ws.column_dimensions[col[0].column_letter].width = max_length + 2

    output_path = WORKSPACE / output_filename
    wb.save(output_path)
    logger.info("Combined %d rows from %d files into %s", len(combined_rows), len(source_files), output_path)

    return _with_download_metadata(
        output_path,
        f"Merged {len(combined_rows)} row(s) from {len(source_files)} file(s) into '{output_filename}'.",
        extra={
            "rows": len(combined_rows),
            "files": len(source_files),
            "columns": columns,
        },
        download_filename=output_filename,
    )


if __name__ == "__main__":
    mcp.run()
